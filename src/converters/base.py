"""
DOCX 파일을 다른 형식으로 변환하는 공통 기능
"""
import os
import json
from io import BytesIO
from typing import Tuple, Dict, Any
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


class DocxExtractor:
    """DOCX 파일에서 텍스트와 구조를 추출하는 공통 클래스"""
    
    def __init__(self, docx_file_path: str = None, docx_content: bytes = None):
        """
        DOCX 파일을 로드합니다.
        
        Args:
            docx_file_path: DOCX 파일 경로
            docx_content: DOCX 파일의 바이트 내용
        """
        if docx_file_path:
            self.doc = DocxDocument(docx_file_path)
        elif docx_content:
            self.doc = DocxDocument(BytesIO(docx_content))
        else:
            raise ValueError("docx_file_path 또는 docx_content 중 하나는 필수입니다.")
    
    def extract_text_content(self) -> str:
        """DOCX에서 순수 텍스트만 추출"""
        text_content = []
        
        # 문서의 모든 단락에서 텍스트 추출
        for paragraph in self.doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # 표의 텍스트도 추출
        for table in self.doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n".join(text_content)
    
    def extract_structured_content(self) -> Dict[str, Any]:
        """DOCX에서 구조화된 데이터 추출"""
        content = {
            "paragraphs": [],
            "tables": [],
            "metadata": {
                "total_paragraphs": 0,
                "total_tables": 0,
                "total_characters": 0
            }
        }
        
        # 단락 추출
        for i, paragraph in enumerate(self.doc.paragraphs):
            if paragraph.text.strip():
                para_data = {
                    "index": i,
                    "text": paragraph.text,
                    "style": paragraph.style.name if paragraph.style else "Normal"
                }
                content["paragraphs"].append(para_data)
        
        # 표 추출
        for i, table in enumerate(self.doc.tables):
            table_data = {
                "index": i,
                "rows": []
            }
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data["rows"].append(row_data)
            
            content["tables"].append(table_data)
        
        # 메타데이터 업데이트
        content["metadata"]["total_paragraphs"] = len(content["paragraphs"])
        content["metadata"]["total_tables"] = len(content["tables"])
        
        full_text = self.extract_text_content()
        content["metadata"]["total_characters"] = len(full_text)
        
        return content


class BaseConverter:
    """변환기의 기본 클래스"""
    
    def __init__(self, extractor: DocxExtractor):
        self.extractor = extractor
    
    def convert(self) -> Tuple[bytes, str]:
        """
        변환을 수행합니다.
        
        Returns:
            Tuple[bytes, str]: (변환된 파일의 바이트 데이터, MIME 타입)
        """
        raise NotImplementedError("하위 클래스에서 구현해야 합니다.")
