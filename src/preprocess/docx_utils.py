"""
DOCX 파싱 유틸리티 모듈
"""
import re
from typing import List, Dict, Any, Tuple, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.shared import Length


class DOCXAnalyzer:
    """DOCX 파일 구조 분석을 위한 유틸리티 클래스"""
    
    @staticmethod
    def analyze_document_structure(doc_path: str) -> Dict[str, Any]:
        """
        DOCX 문서의 구조를 분석하여 메타데이터 반환
        """
        try:
            doc = Document(doc_path)
            
            analysis = {
                "total_paragraphs": len(doc.paragraphs),
                "total_tables": len(doc.tables),
                "styles_used": set(),
                "heading_levels": set(),
                "text_patterns": {
                    "titles": [],
                    "headings": [],
                    "sections": [],
                    "lists": [],
                    "urls": []
                },
                "document_stats": {
                    "word_count": 0,
                    "character_count": 0,
                    "table_cells": 0
                }
            }
            
            # 단락 분석
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # 스타일 수집
                if para.style:
                    analysis["styles_used"].add(para.style.name)
                    
                    # 헤딩 레벨 추출
                    if "heading" in para.style.name.lower():
                        heading_match = re.search(r'heading (\d+)', para.style.name.lower())
                        if heading_match:
                            analysis["heading_levels"].add(int(heading_match.group(1)))
                
                # 패턴 분석
                DOCXAnalyzer._analyze_text_patterns(text, analysis["text_patterns"])
                
                # 통계
                analysis["document_stats"]["word_count"] += len(text.split())
                analysis["document_stats"]["character_count"] += len(text)
            
            # 표 분석
            for table in doc.tables:
                cell_count = sum(len(row.cells) for row in table.rows)
                analysis["document_stats"]["table_cells"] += cell_count
            
            # 세트를 리스트로 변환 (JSON 직렬화를 위해)
            analysis["styles_used"] = list(analysis["styles_used"])
            analysis["heading_levels"] = list(analysis["heading_levels"])
            
            return analysis
            
        except Exception as e:
            return {"error": str(e)}
    
    @staticmethod
    def _analyze_text_patterns(text: str, patterns: Dict[str, List[str]]):
        """텍스트에서 패턴을 분석하여 분류"""
        
        # 제목 패턴
        title_patterns = [
            r'^\[([^\]]+)\]$',
            r'^【([^】]+)】$',
            r'^\*\*([^*]+)\*\*$',
            r'^◎\s*(.+)$',
            r'^●\s*(.+)$'
        ]
        
        # 헤딩 패턴
        heading_patterns = [
            r'^(\d+\.?\s*.+)$',
            r'^([가-힣]+\s*\d+\.?\s*.+)$',
            r'^(\w+\.\s*.+)$',
            r'^([IVX]+\.\s*.+)$',
            r'^(제\d+[조항절]\s*.+)$'
        ]
        
        # 섹션 패턴
        section_patterns = [
            r'^(\d+-\d+\)\s*.+)$',
            r'^(\d+\)\s*.+)$',
            r'^([①-⑳]\s*.+)$',
            r'^([가-힣]\)\s*.+)$',
            r'^([ㄱ-ㅎ]\.\s*.+)$'
        ]
        
        # URL 패턴
        url_pattern = r'https?://[^\s]+'
        
        # 패턴 매칭
        for pattern in title_patterns:
            if re.match(pattern, text):
                patterns["titles"].append(text)
                return
        
        for pattern in heading_patterns:
            if re.match(pattern, text):
                patterns["headings"].append(text)
                return
        
        for pattern in section_patterns:
            if re.match(pattern, text):
                patterns["sections"].append(text)
                return
        
        # URL 찾기
        urls = re.findall(url_pattern, text)
        if urls:
            patterns["urls"].extend(urls)
        
        # 리스트 패턴 (-, *, 숫자. 등으로 시작)
        if re.match(r'^[\-\*\+]\s+', text) or re.match(r'^\d+\.\s+', text):
            patterns["lists"].append(text)


class ContentExtractor:
    """DOCX에서 컨텐츠를 추출하는 클래스"""
    
    def __init__(self):
        self.extracted_content = []
    
    def extract_all_content(self, doc_path: str) -> List[Dict[str, Any]]:
        """문서의 모든 컨텐츠를 순서대로 추출"""
        try:
            doc = Document(doc_path)
            content_items = []
            
            # 단락들 추출
            for paragraph in doc.paragraphs:
                para_content = self._extract_paragraph_content(paragraph)
                if para_content:
                    content_items.append(para_content)
            
            # 표들 추출
            for table in doc.tables:
                table_content = self._extract_table_content(table)
                if table_content:
                    content_items.append(table_content)
            
            return content_items
            
        except Exception as e:
            return [{"type": "error", "content": str(e)}]
    
    def _extract_paragraph_content(self, paragraph: Paragraph) -> Optional[Dict[str, Any]]:
        """단락에서 컨텐츠 추출"""
        text = paragraph.text.strip()
        if not text:
            return None
        
        # 스타일 정보 추출
        style_info = {
            "style_name": paragraph.style.name if paragraph.style else "",
            "is_bold": any(run.bold for run in paragraph.runs if run.bold),
            "font_sizes": [run.font.size.pt for run in paragraph.runs if run.font.size],
            "alignment": str(paragraph.alignment) if paragraph.alignment else ""
        }
        
        return {
            "type": "paragraph",
            "content": text,
            "style": style_info
        }
    
    def _extract_table_content(self, table: Table) -> Dict[str, Any]:
        """표에서 컨텐츠 추출"""
        rows_data = []
        
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                row_data.append({
                    "text": cell_text,
                    "row": row_idx,
                    "col": cell_idx
                })
            rows_data.append(row_data)
        
        return {
            "type": "table",
            "content": rows_data,
            "dimensions": {
                "rows": len(table.rows),
                "columns": len(table.columns) if table.rows else 0
            }
        }


class PatternMatcher:
    """텍스트 패턴 매칭을 위한 클래스"""
    
    def __init__(self):
        self.patterns = {
            "title": [
                r'^\[([^\]]+)\]$',           # [제목]
                r'^【([^】]+)】$',             # 【제목】
                r'^\*\*([^*]+)\*\*$',        # **제목**
                r'^◎\s*(.+)$',              # ◎ 제목
                r'^●\s*(.+)$',              # ● 제목
                r'^#+\s*(.+)$',             # # 제목 (마크다운)
            ],
            "heading": [
                r'^(\d+\.\s*.+)$',           # 1. 제목 (숫자 + 점 + 공백)
                r'^([가-힣]+\s*\d+\.?\s*.+)$', # 한글 + 숫자
                r'^(\w+\.\s*.+)$',           # 알파벳. 제목
                r'^([IVX]+\.\s*.+)$',        # 로마숫자. 제목
                r'^(제\d+[조항절]\s*.+)$',    # 제1조 형태
                r'^(Chapter\s+\d+.*)$',      # Chapter 1
                r'^(Section\s+\d+.*)$',      # Section 1
            ],
            "section": [
                r'^(\d+-\d+\)\s*.+)$',       # 1-1) (숫자-숫자))
                r'^(\d+\)\s*.+)$',           # 1) (숫자))
                r'^([①-⑳]\s*.+)$',          # ① (동그라미 숫자)
                r'^([가-힣]\)\s*.+)$',        # 가) (한글))
                r'^([ㄱ-ㅎ]\.\s*.+)$',        # ㄱ. (자음.)
                r'^([a-z]\)\s*.+)$',         # a) (소문자))
                r'^([A-Z]\.\s*.+)$',         # A. (대문자.)
            ],
            "list": [
                r'^[\-\*\+]\s+(.+)$',        # - 항목
                r'^(\d+)\.\s+(.+)$',         # 1. 항목 (하지만 heading과 겹칠 수 있음)
                r'^•\s*(.+)$',               # • 항목
                r'^◦\s*(.+)$',               # ◦ 항목
            ],
            "url": [
                r'(https?://[^\s]+)',        # HTTP URL
                r'(www\.[^\s]+)',            # www URL
                r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', # 이메일
            ],
            "special": [
                r'^\([^\)]+\)$',             # (괄호 안 내용)
                r'^\[[^\]]+\]\([^\)]+\)$',   # [링크](URL) 마크다운
                r'^>+\s*(.+)$',              # > 인용문
                r'^`{1,3}(.+)`{1,3}$',       # 코드 블록
            ]
        }
    
    def match_pattern(self, text: str) -> Tuple[str, Optional[str]]:
        """
        텍스트가 어떤 패턴에 해당하는지 확인
        
        Returns:
            Tuple[str, Optional[str]]: (패턴_타입, 매칭된_그룹)
        """
        text = text.strip()
        
        for pattern_type, pattern_list in self.patterns.items():
            for pattern in pattern_list:
                match = re.match(pattern, text, re.MULTILINE | re.DOTALL)
                if match:
                    matched_group = match.group(1) if match.groups() else match.group(0)
                    return pattern_type, matched_group
        
        return "text", None
    
    def extract_all_patterns(self, text: str) -> Dict[str, List[str]]:
        """텍스트에서 모든 패턴을 추출"""
        results = {pattern_type: [] for pattern_type in self.patterns.keys()}
        
        lines = text.split('\n')
        for line in lines:
            pattern_type, matched = self.match_pattern(line)
            if matched:
                results[pattern_type].append(matched)
        
        return results


def validate_json_structure(data: List[Dict[str, Any]]) -> Tuple[bool, List[str]]:
    """
    생성된 JSON 구조가 index.json 형태와 일치하는지 검증
    
    Returns:
        Tuple[bool, List[str]]: (유효성, 오류_목록)
    """
    errors = []
    
    if not isinstance(data, list):
        errors.append("최상위 구조가 리스트가 아닙니다")
        return False, errors
    
    for i, item in enumerate(data):
        if not isinstance(item, dict):
            errors.append(f"항목 {i}가 딕셔너리가 아닙니다")
            continue
        
        # 필수 필드 확인
        if "filename" not in item:
            errors.append(f"항목 {i}에 'filename' 필드가 없습니다")
        
        if "data" not in item:
            errors.append(f"항목 {i}에 'data' 필드가 없습니다")
            continue
        
        if not isinstance(item["data"], list):
            errors.append(f"항목 {i}의 'data' 필드가 리스트가 아닙니다")
            continue
        
        # data 항목들 검증
        for j, data_item in enumerate(item["data"]):
            if not isinstance(data_item, dict):
                errors.append(f"항목 {i}의 data[{j}]가 딕셔너리가 아닙니다")
                continue
            
            if "content" not in data_item:
                errors.append(f"항목 {i}의 data[{j}]에 'content' 필드가 없습니다")
    
    return len(errors) == 0, errors
