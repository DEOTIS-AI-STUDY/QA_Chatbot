import os
from docx import Document
from docx.oxml import ns

def extract_hyperlinks(paragraph):
    nsmap = ns.nsmap
    links = []
    for elem in paragraph._element.findall('.//w:hyperlink', nsmap):
        rId = elem.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
        url = None
        text = ""
        if rId:
            rel = paragraph.part.rels.get(rId)
            if rel:
                url = rel.target_ref
        for t in elem.findall('.//w:t', nsmap):
            text += t.text or ""
        if url:
            links.append((text, url))
    return links

def get_cell_text(cell):
    return "\n".join([p.text for p in cell.paragraphs]).strip()

def parse_table(table):
    lines = []
    rows = table.rows
    if len(rows) < 2:
        return lines
    header = [get_cell_text(cell).replace('\n', ' ') for cell in rows[0].cells]
    md_header = "| " + " | ".join(header) + " |"
    md_sep = "| " + " | ".join(["---"] * len(header)) + " |"
    lines.append(md_header)
    lines.append(md_sep)
    for row in rows[1:]:
        row_items = []
        for col_idx, cell in enumerate(row.cells):
            if col_idx < len(header):
                value = get_cell_text(cell).replace('\n', ' ')
                row_items.append(value)
        lines.append("| " + " | ".join(row_items) + " |")
    lines.append("")
    lines.append("")
    return lines

def parse_docx_file(filepath):
    doc = Document(filepath)
    nsmap = ns.nsmap
    result_lines = []
    tables_blocks = []
    img_count = 1

    block_items = []
    for block in doc.element.body:
        if block.tag.endswith('tbl'):
            for table in doc.tables:
                if table._tbl == block:
                    block_items.append(('table', table))
                    break
        elif block.tag.endswith('p'):
            for para in doc.paragraphs:
                if para._p == block:
                    block_items.append(('para', para))
                    break

    skip_table_indices = set()
    # 표와 표 위 2단락 추출
    for idx, (kind, item) in enumerate(block_items):
        if kind == 'table':
            # 표 위 2단락 추출
            prev_paras = []
            for i in range(max(0, idx-2), idx):
                if block_items[i][0] == 'para':
                    prev_paras.append(block_items[i][1].text.strip())
            table = item
            table_lines = parse_table(table)
            # 괄호로 묶어서 저장
            block = "{\n"
            for para_text in prev_paras:
                if para_text:
                    block += para_text + "\n"
            block += "\n".join(table_lines)
            block += "\n}\n"
            tables_blocks.append(block)
            # 표 위 2단락은 data.txt에 남기고, 표는 제거
            skip_table_indices.add(idx)
        elif kind == 'para':
            para = item
            text = para.text.strip()
            if text:
                result_lines.append(text)
            links = extract_hyperlinks(para)
            if links:
                for link_text, url in links:
                    if link_text.strip().lower().startswith("http://") or link_text.strip().lower().startswith("https://"):
                        continue
                    else:
                        result_lines.append(f"{link_text}({url})")
            for drawing in para._element.findall('.//w:drawing', nsmap):
                docpr = drawing.find('.//wp:docPr', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
                img_name = docpr.attrib.get('name') if docpr is not None else f"그림 {img_count}"
                img_descr = docpr.attrib.get('descr') if docpr is not None else ""
                hlink = drawing.find('.//a:hlinkClick', {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                url = ""
                if hlink is not None:
                    rId = hlink.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    if rId:
                        rel = doc.part.rels.get(rId)
                        if rel:
                            url = rel.target_ref
                if url:
                    result_lines.append(f"이미지:{img_name}, desc:{img_descr}, url:{url}")
                else:
                    result_lines.append(f"이미지:{img_name}, desc:{img_descr}")
                img_count += 1

    # 표 바로 위 2단락은 data.txt에 남기고, 표는 제거
    # 이미 result_lines에 표 위 2단락이 들어가 있으므로, 표 내용만 data.txt에서 빠짐

    return "\n".join(result_lines), tables_blocks

def main():

    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../'))
    data_dir = os.path.join(base_dir, 'data')
    output_dir = os.path.join(data_dir, 'docx')
    print(f"[INFO] 변환 대상 폴더: {data_dir}")
    print(f"[INFO] 결과 저장 폴더: {output_dir}")
    os.makedirs(output_dir, exist_ok=True)

    files = [f for f in os.listdir(data_dir) if os.path.isfile(os.path.join(data_dir, f))]
    print(f"[INFO] data 폴더 내 파일 목록: {files}")
    docx_files = [f for f in files if f.lower().endswith('.docx')]
    print(f"[INFO] 변환할 docx 파일 목록: {docx_files}")

    for filename in docx_files:
        print(f"[INFO] 파일 처리 시작: {filename}")
        filepath = os.path.join(data_dir, filename)
        try:
            doc = Document(filepath)
        except Exception as e:
            print(f"[ERROR] 파일 열기 실패: {filename}, 에러: {e}")
            continue
        nsmap = ns.nsmap

        # 새 문서 생성
        new_doc = Document()

        img_count = 1
        for block in doc.element.body:
            if block.tag.endswith('tbl'):
                # 표를 마크다운 텍스트로 변환 후 기존 표는 삭제
                for table in doc.tables:
                    if table._tbl == block:
                        print(f"[INFO] 표 발견 및 변환: {filename}")
                        md_lines = parse_table(table)
                        if md_lines:
                            for md_line in md_lines:
                                new_doc.add_paragraph(md_line)
                        break
            elif block.tag.endswith('p'):
                for para in doc.paragraphs:
                    if para._p == block:
                        # 하이퍼링크 변환
                        links = extract_hyperlinks(para)
                        if links:
                            for link_text, url in links:
                                print(f"[INFO] 하이퍼링크 변환: {link_text}({url}) in {filename}")
                                if link_text.strip():
                                    new_doc.add_paragraph(f"{link_text}({url})")
                        # 이미지 처리
                        nsmap = ns.nsmap
                        img_found = False
                        for drawing in para._element.findall('.//w:drawing', nsmap):
                            img_found = True
                            docpr = drawing.find('.//wp:docPr', {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
                            img_name = docpr.attrib.get('name') if docpr is not None else f"그림 {img_count}"
                            img_descr = docpr.attrib.get('descr') if docpr is not None else ""
                            hlink = drawing.find('.//a:hlinkClick', {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                            url = ""
                            if hlink is not None:
                                rId = hlink.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                                if rId:
                                    rel = doc.part.rels.get(rId)
                                    if rel:
                                        url = rel.target_ref
                            img_info = f"이미지: {img_name}, desc: {img_descr}"
                            if url:
                                img_info += f", url: {url}"
                            print(f"[INFO] 이미지 변환: {img_info} in {filename}")
                            new_doc.add_paragraph(img_info)
                            img_count += 1
                        # 테이블/하이퍼링크/이미지에 해당하지 않는 경우 일반 텍스트만 추가
                        if not links and not img_found:
                            if para.text.strip():
                                new_doc.add_paragraph(para.text.strip())
                        break

        # 저장
        out_path = os.path.join(output_dir, filename)
        try:
            new_doc.save(out_path)
            print(f"[SUCCESS] 변환된 파일 저장: {out_path}")
        except Exception as e:
            print(f"[ERROR] 파일 저장 실패: {out_path}, 에러: {e}")

if __name__ == "__main__":
    main()

class DocParser:
    def __init__(self, filepath):
        self.filepath = filepath
        self.data_txt = None
        self.tables_txt = None
        self._parse_file()

    def _parse_file(self):
        # 기존 parse_docx_file 함수 활용
        text, tables_blocks = parse_docx_file(self.filepath)
        self.data_txt = text
        self.tables_txt = tables_blocks

    def get_text(self):
        return self.data_txt

    def get_tables(self):
        # tables_txt는 이미 {}로 묶인 블록 리스트임
        result = []
        for block in self.tables_txt:
            # {} 안 내용만 추출
            start = block.find('{')
            end = block.rfind('}')
            if start != -1 and end != -1:
                content = block[start+1:end].strip()
                result.append(content)
        return result