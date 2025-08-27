"""
Elasticsearch 인덱싱 전용 유틸리티
"""
import os
import json
import re
from typing import List, Tuple, Any
from elasticsearch import Elasticsearch
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.vectorstores import ElasticsearchStore
from langchain.schema import Document
from docx import Document as DocxDocument
from core.config import ELASTICSEARCH_URL, INDEX_NAME
from .elasticsearch import ElasticsearchManager, get_optimized_text_splitter


class ElasticsearchIndexer:
    """Elasticsearch 인덱싱 전용 클래스"""
    
    @staticmethod
    def delete_documents_by_category(categories: List[str]) -> Tuple[bool, str]:
        """특정 카테고리의 문서들을 선택적으로 삭제"""
        try:
            from utils.elasticsearch import ElasticsearchManager
            config = ElasticsearchManager.get_connection_config()
            es = Elasticsearch(**config)
            
            if not es.indices.exists(index=INDEX_NAME):
                return True, "인덱스가 존재하지 않습니다."
            
            # 삭제할 카테고리들에 대한 쿼리 생성
            delete_query = {
                "query": {
                    "terms": {
                        "metadata.category.keyword": categories
                    }
                }
            }
            
            # 삭제 전 문서 수 확인
            count_before = es.count(index=INDEX_NAME).get("count", 0)
            
            # 선택적 삭제 실행
            delete_response = es.delete_by_query(
                index=INDEX_NAME,
                body=delete_query,
                refresh=True
            )
            
            deleted_count = delete_response.get("deleted", 0)
            count_after = es.count(index=INDEX_NAME).get("count", 0)
            
            return True, f"카테고리 {categories} 문서 {deleted_count}개 삭제됨 (전체: {count_before} → {count_after})"
            
        except Exception as e:
            return False, f"선택적 삭제 오류: {str(e)}"
    
    @staticmethod
    def get_existing_files_by_category(category: str) -> List[str]:
        """특정 카테고리의 기존 파일명 목록 반환"""
        try:
            from utils.elasticsearch import ElasticsearchManager
            config = ElasticsearchManager.get_connection_config()
            es = Elasticsearch(**config)
            
            if not es.indices.exists(index=INDEX_NAME):
                return []
            
            # 카테고리별 파일명 집계 쿼리
            agg_query = {
                "size": 0,
                "query": {
                    "term": {
                        "metadata.category.keyword": category
                    }
                },
                "aggs": {
                    "filenames": {
                        "terms": {
                            "field": "metadata.filename.keyword",
                            "size": 1000
                        }
                    }
                }
            }
            
            response = es.search(index=INDEX_NAME, body=agg_query)
            buckets = response.get("aggregations", {}).get("filenames", {}).get("buckets", [])
            
            return [bucket["key"] for bucket in buckets]
            
        except Exception as e:
            print(f"기존 파일 목록 조회 오류: {e}")
            return []
    
    # 파일 목록 조회 함수들 (인덱싱에서 사용)
    @staticmethod
    def list_pdfs(pdf_dir: str) -> List[str]:
        """PDF 파일 목록 반환"""
        try:
            names = os.listdir(pdf_dir)
        except FileNotFoundError:
            return []
        
        files = []
        for name in names:
            if name.lower().endswith(".pdf"):
                files.append(os.path.join(pdf_dir, name))
        return sorted(files)
    
    @staticmethod
    def list_txt_files(txt_dir: str) -> List[str]:
        """TXT 파일 목록 반환"""
        try:
            names = os.listdir(txt_dir)
        except FileNotFoundError:
            return []
        
        files = []
        for name in names:
            if name.lower().endswith(".txt"):
                files.append(os.path.join(txt_dir, name))
        return sorted(files)
    
    @staticmethod
    def list_docx_files(docx_dir: str) -> List[str]:
        """DOCX 파일 목록 반환"""
        try:
            names = os.listdir(docx_dir)
        except FileNotFoundError:
            return []
        
        files = []
        for name in names:
            if name.lower().endswith((".docx", ".doc")):
                files.append(os.path.join(docx_dir, name))
        return sorted(files)
    
    @staticmethod
    def list_json_files(json_dir: str) -> List[str]:
        """JSON 파일 목록 반환"""
        try:
            names = os.listdir(json_dir)
        except FileNotFoundError:
            return []
        
        files = []
        for name in names:
            if name.lower().endswith(".json"):
                files.append(os.path.join(json_dir, name))
        return sorted(files)
    
    @staticmethod
    def index_pdfs(pdf_files: List[str], embeddings, hybrid_tracker, incremental: bool = True) -> Tuple[bool, str]:
        """PDF 파일들을 Elasticsearch에 인덱싱"""
        return ElasticsearchIndexer._index_files_by_type(
            files=pdf_files,
            file_type="PDF",
            embeddings=embeddings,
            hybrid_tracker=hybrid_tracker,
            incremental=incremental,
            loader_func=lambda path: PyPDFLoader(path).load()
        )
    
    @staticmethod
    def index_txt_files(txt_files: List[str], embeddings, hybrid_tracker, incremental: bool = True) -> Tuple[bool, str]:
        """TXT 파일들을 Elasticsearch에 인덱싱"""
        return ElasticsearchIndexer._index_files_by_type(
            files=txt_files,
            file_type="TXT",
            embeddings=embeddings,
            hybrid_tracker=hybrid_tracker,
            incremental=incremental,
            loader_func=lambda path: TextLoader(path, encoding='utf-8').load()
        )
    
    @staticmethod
    def index_json_files(json_files: List[str], embeddings, hybrid_tracker, incremental: bool = True) -> Tuple[bool, str]:
        """JSON 파일들을 Elasticsearch에 인덱싱"""
        def json_loader(path):
            with open(path, 'r', encoding='utf-8') as f:
                json_data = json.load(f)
            text_content = ElasticsearchIndexer._extract_text_from_json(json_data)
            if text_content:
                return [Document(
                    page_content=text_content,
                    metadata={"source": path, "filename": os.path.basename(path), "category": "JSON"}
                )]
            return []
        
        return ElasticsearchIndexer._index_files_by_type(
            files=json_files,
            file_type="JSON",
            embeddings=embeddings,
            hybrid_tracker=hybrid_tracker,
            incremental=incremental,
            loader_func=json_loader
        )
    
    @staticmethod
    def index_docx_files(docx_files: List[str], embeddings, hybrid_tracker, incremental: bool = True) -> Tuple[bool, str]:
        """DOCX 파일들을 Elasticsearch에 인덱싱"""
        def docx_loader(path):
            try:
                # DocParser 사용 (있는 경우)
                from utils.docparser import DocParser
                parser = DocParser(path)
                text_content = parser.get_text()
                
                docs = [Document(
                    page_content=text_content,
                    metadata={"source": path, "filename": os.path.basename(path), "category": "DOCX"}
                )]
                
                # 표 내용 추가
                table_blocks = parser.get_tables()
                for table_content in table_blocks:
                    docs.append(Document(
                        page_content=table_content,
                        metadata={"source": path, "filename": os.path.basename(path), "category": "DOCX"}
                    ))
                
                return docs
            except ImportError:
                # DocParser가 없으면 기본 python-docx 사용
                return ElasticsearchIndexer._docx_basic_loader(path)
        
        return ElasticsearchIndexer._index_files_by_type(
            files=docx_files,
            file_type="DOCX",
            embeddings=embeddings,
            hybrid_tracker=hybrid_tracker,
            incremental=incremental,
            loader_func=docx_loader
        )
    
    @staticmethod
    def _docx_basic_loader(path: str) -> List[Document]:
        """기본 DOCX 로더 (python-docx 사용)"""
        doc = DocxDocument(path)
        text_content = []
        
        # 단락 텍스트 추출
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # 표 텍스트 추출
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        if text_content:
            full_text = "\n".join(text_content)
            return [Document(
                page_content=full_text,
                metadata={"source": path, "filename": os.path.basename(path), "category": "DOCX"}
            )]
        return []

    @staticmethod
    def chunk_pdf_with_md_tables(pages: Document, text_splitter) -> list[Document]:
        # pages : loaded datas 
        # file_path, file_type : for metadata

        # 1단계: 표와 텍스트 덩어리를 분리 (향상된 로직)
        preliminary_chunks = []
        current_text_buffer = ""
        current_table_buffer = ""
        in_table = False
        metadata_buffer = {}

        for page in pages:
            lines = page.page_content.split('\n')
            for line in lines:
                stripped_line = line.strip()
                if not stripped_line:
                    continue

                # Case 1: 줄이 '|'로 시작하는 경우 (표의 시작 또는 새 행)
                if stripped_line.startswith('|'):
                    if not in_table:
                        # 표가 새로 시작되면, 이전까지의 텍스트를 청크로 저장
                        if current_text_buffer.strip():
                            preliminary_chunks.append(Document(page_content=current_text_buffer.strip(), metadata=metadata_buffer))
                        current_text_buffer = ""
                        in_table = True
                        metadata_buffer = page.metadata.copy()
                    
                    current_table_buffer += stripped_line + "\n"
                
                # Case 2: '|'로 시작하지 않지만, 현재 표 내부에 있는 경우
                elif in_table:
                    # 테이블 버퍼의 마지막 줄이 '|'로 끝나지 않았다면, 현재 줄은 그 행의 연속임
                    last_line_in_buffer = current_table_buffer.rstrip('\n').split('\n')[-1]
                    if last_line_in_buffer and not last_line_in_buffer.endswith('|'):
                        # 이전 행의 끝(개행문자)을 지우고, 공백과 함께 현재 내용을 이어붙임
                        current_table_buffer = current_table_buffer.rstrip('\n') + " " + stripped_line + "\n"
                    else:
                        # 이전 행이 완성된 행이었다면, 표가 끝난 것으로 간주
                        if current_table_buffer.strip():
                            preliminary_chunks.append(Document(page_content=current_table_buffer.strip(), metadata=metadata_buffer))
                        current_table_buffer = ""
                        in_table = False
                        metadata_buffer = page.metadata.copy()
                        current_text_buffer += line + "\n"
                
                # Case 3: 표의 일부가 아닌 일반 텍스트
                else:
                    if not metadata_buffer or metadata_buffer.get('page') != page.metadata.get('page'):
                        metadata_buffer = page.metadata.copy()
                    current_text_buffer += line + "\n"

        # 루프 종료 후 남아있는 버퍼 처리
        if current_table_buffer.strip():
            preliminary_chunks.append(Document(page_content=current_table_buffer.strip(), metadata=metadata_buffer))
        if current_text_buffer.strip():
            preliminary_chunks.append(Document(page_content=current_text_buffer.strip(), metadata=metadata_buffer))

        # 2단계: 텍스트 덩어리만 추가 분할
        final_chunks = []

        exchunk = ""
        for chunk in preliminary_chunks:
            # chunk.metadata.update({
            #                 "source": file_path,
            #                 "filename": os.path.basename(file_path),
            #                 "category": file_type
            #             })
            if chunk.page_content.strip().startswith('|'):
                if exchunk != "":
                    last_line = exchunk.page_content.strip().splitlines()[-1]
                    chunk.page_content = last_line+'\n' + chunk.page_content
                final_chunks.append(chunk) # 표는 그대로 추가
            else:
                # 텍스트는 잘게 분할하여 추가
                sub_chunks = text_splitter.create_documents([chunk.page_content], metadatas=[chunk.metadata])
                final_chunks.extend(sub_chunks)
            exchunk = chunk
                
        return final_chunks
    
    @staticmethod
    def _index_files_by_type(files: List[str], file_type: str, embeddings, hybrid_tracker, 
                           incremental: bool, loader_func) -> Tuple[bool, str]:
        """파일 타입별 통합 인덱싱 로직"""
        stage_name = f"{file_type}_인덱싱_시작"
        hybrid_tracker.track_preprocessing_stage(stage_name)
        
        # 증분 업데이트 처리
        if not incremental:
            ElasticsearchIndexer._delete_entire_index()
        else:
            success, message = ElasticsearchIndexer.delete_documents_by_category([file_type])
            if not success:
                hybrid_tracker.end_preprocessing_stage(stage_name)
                return False, f"기존 {file_type} 문서 삭제 실패: {message}"
            print(f"📚 {message}")
        
        all_documents = []
        
        # 파일별 처리
        for file_path in files:
            file_stage = f"{file_type}_처리_{os.path.basename(file_path)}"
            hybrid_tracker.track_preprocessing_stage(file_stage)
            
            try:
                docs = loader_func(file_path)
                # this docs will be parameter in chunker

                splitter = get_optimized_text_splitter()
                if file_type.upper() == "PDF":
                    chunks = ElasticsearchIndexer.chunk_pdf_with_md_tables(docs, splitter)
                    all_documents.extend(chunks)
                else:
                    for doc in docs:
                        # if "표" in doc.page_content or len(doc.page_content.split('\n')) > 5:
                        #     # 표이거나 여러 줄인 경우 분할하지 않음
                        #     all_documents.append(doc)
                        # else:
                        # 일반 텍스트는 분할
                        chunks = splitter.split_documents([doc])
                        all_documents.extend(chunks)
                    # 메타데이터 보강
                for doc in all_documents[-len(docs):]:  # 방금 추가된 문서들만
                    doc.metadata.update({
                        "source": file_path,
                        "filename": os.path.basename(file_path),
                        "category": file_type
                    })
                
                
                
            except Exception as e:
                print(f"{file_type} 파일 처리 오류 ({file_path}): {e}")
                continue
            finally:
                hybrid_tracker.end_preprocessing_stage(file_stage)
        
        if not all_documents:
            hybrid_tracker.end_preprocessing_stage(stage_name)
            return False, "추출된 텍스트가 없습니다."
        
        # PDF 파일의 경우 처리 결과를 export
        if file_type.upper() == "PDF":
            export_path = ElasticsearchIndexer._export_pdf_processing_data(all_documents, "data")
            if export_path:
                print(f"📊 PDF 처리 결과가 export되었습니다: {export_path}")
        
        # Elasticsearch에 저장
        return ElasticsearchIndexer._save_to_elasticsearch(all_documents, embeddings, hybrid_tracker, stage_name, file_type)
    
    @staticmethod
    def _delete_entire_index():
        """전체 인덱스 삭제"""
        config = ElasticsearchManager.get_connection_config()
        es = Elasticsearch(**config)
        if es.indices.exists(index=INDEX_NAME):
            es.indices.delete(index=INDEX_NAME)
    
    @staticmethod
    def _save_to_elasticsearch(documents: List[Document], embeddings, hybrid_tracker, stage_name: str, file_type: str) -> Tuple[bool, str]:
        """Elasticsearch에 문서 저장"""
        save_stage = "Elasticsearch_저장"
        hybrid_tracker.track_preprocessing_stage(save_stage)
        
        try:
            es_client, success, message = ElasticsearchManager.get_safe_elasticsearch_client()
            if not success:
                raise Exception(f"Elasticsearch 연결 실패: {message}")
            
            mapping = {
                "mappings":{
                    "properties":{
                        "text":{
                            "type":"text",
                            "analyzer":"nori"
                        }
                    }
                }
            }

            if es_client.indices.exists(index=INDEX_NAME):
                es_client.indices.delete(index=INDEX_NAME)
            es_client.indices.create(index=INDEX_NAME, body=mapping)

            ElasticsearchStore.from_documents(
                documents,
                embedding=embeddings,
                es_url=ELASTICSEARCH_URL,
                index_name=INDEX_NAME
            )
            
            es_client.indices.refresh(index=INDEX_NAME)
            cnt = es_client.count(index=INDEX_NAME).get("count", 0)
            
            return True, f"{file_type} 인덱싱 완료. 문서 수: {cnt}"
            
        except Exception as e:
            return False, f"{file_type} 인덱싱 오류: {str(e)}"
        finally:
            hybrid_tracker.end_preprocessing_stage(save_stage)
            hybrid_tracker.end_preprocessing_stage(stage_name)
    
    @staticmethod
    def _export_pdf_processing_data(data, export_dir: str = "data") -> str:
        """PDF 처리 결과를 JSON으로 export하는 함수 - 다양한 데이터 형태 지원"""
        try:
            import uuid
            from datetime import datetime
            
            # export 디렉토리 생성
            os.makedirs(export_dir, exist_ok=True)
            
            # export 파일명 생성
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            export_filename = f"pdf_indexing_export_{timestamp}.json"
            export_path = os.path.join(export_dir, export_filename)
            
            # 데이터 형태 판별 및 처리
            processed_documents = []
            
            if isinstance(data, list) and data:
                # 첫 번째 요소로 데이터 형태 판별
                first_item = data[0]
                
                if hasattr(first_item, 'page_content'):
                    # Document 객체 리스트인 경우 (_index_files_by_type에서 호출)
                    processed_documents = data
                elif isinstance(first_item, dict) and 'content' in first_item:
                    # 딕셔너리 리스트인 경우 (index_all_files에서 호출)
                    for item in data:
                        # 딕셔너리를 Document 형태로 변환
                        content = item.get('content', '')
                        metadata = item.get('metadata', {})
                        
                        # Document-like 객체 생성 (page_content와 metadata 속성을 가진 객체)
                        class DocumentLike:
                            def __init__(self, page_content, metadata):
                                self.page_content = page_content
                                self.metadata = metadata
                        
                        processed_documents.append(DocumentLike(content, metadata))
                else:
                    raise ValueError(f"지원하지 않는 데이터 형태입니다: {type(first_item)}")
            else:
                print("❌ 빈 데이터 또는 잘못된 형태입니다.")
                return ""
            
            # 문서 분석
            export_data = {
                "export_info": {
                    "export_time": datetime.now().isoformat(),
                    "total_documents": len(processed_documents),
                    "file_type": "PDF",
                    "description": "PDF 인덱싱 과정에서 vector DB에 저장될 문서들"
                },
                "documents": []
            }
            
            # 표 관련 통계
            table_documents = 0
            mixed_documents = 0
            total_table_markers = 0
            total_markdown_lines = 0
            
            for doc in processed_documents:
                # 문서 ID 생성
                doc_id = str(uuid.uuid4())
                
                # 표 관련 분석
                content = doc.page_content
                has_table = ElasticsearchIndexer._is_table_content(content)
                table_markers = content.count('**[표')
                markdown_table_lines = len([line for line in content.split('\n') if '|' in line and line.strip()])
                
                if has_table:
                    table_documents += 1
                if table_markers > 0:
                    mixed_documents += 1
                
                total_table_markers += table_markers
                total_markdown_lines += markdown_table_lines
                
                # 문서 정보 구성
                doc_info = {
                    "document_id": doc_id,
                    "filename": doc.metadata.get('filename', 'unknown'),
                    "source": doc.metadata.get('source', ''),
                    "content_length": len(content),
                    "line_count": len(content.split('\n')),
                    "has_table": has_table,
                    "table_markers_found": table_markers,
                    "markdown_table_lines": markdown_table_lines,
                    "metadata": doc.metadata,
                    "content": content
                }
                
                # 페이지 정보가 있으면 추가
                if 'page' in doc.metadata:
                    doc_info["page"] = doc.metadata['page']
                
                # 표 개수 정보가 있으면 추가
                if 'table_count' in doc.metadata:
                    doc_info["table_count"] = doc.metadata['table_count']
                
                export_data["documents"].append(doc_info)
            
            # 통계 정보 업데이트
            export_data["export_info"].update({
                "total_table_documents": table_documents,
                "total_mixed_documents": mixed_documents,
                "table_detection_ratio": f"{table_documents}/{len(processed_documents)} ({table_documents/len(processed_documents)*100:.1f}%)" if processed_documents else "0/0 (0%)"
            })
            
            # JSON 파일로 저장
            with open(export_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, ensure_ascii=False, indent=2)
            
            print(f"📊 PDF 처리 데이터 export 완료: {export_path}")
            print(f"   - 총 문서: {len(processed_documents)}개")
            print(f"   - 표 포함 문서: {table_documents}개")
            print(f"   - 인라인 표 마커: {total_table_markers}개")
            print(f"   - 마크다운 표 라인: {total_markdown_lines}개")
            
            return export_path
            
        except Exception as e:
            print(f"❌ PDF 데이터 export 오류: {e}")
            import traceback
            print(f"오류 상세: {traceback.format_exc()}")
            return ""
    
    @staticmethod
    def _is_table_content(content: str) -> bool:
        """표 내용인지 판단하는 함수 - 개선된 버전"""
        if not content or len(content.strip()) < 10:
            return False
        
        # 1. 인라인 표 마커가 있는 경우 (**[표 N]**)
        if re.search(r'\*\*\[표\s*\d*\]\*\*', content):
            return True
        
        # 2. mixed 타입 문서인지 확인 (메타데이터에서)
        if 'type' in content and 'mixed' in content:
            return True
        
        # 3. 마크다운 표 형식이 포함된 경우
        lines = content.split('\n')
        pipe_lines = [line for line in lines if '|' in line and line.strip()]
        
        if len(pipe_lines) >= 3:  # 최소 3줄 이상
            # 헤더와 구분선이 있는지 확인
            has_separator = any('---' in line or '===' in line for line in pipe_lines[:5])
            if has_separator:
                # 파이프 라인 비율 확인
                pipe_ratio = len(pipe_lines) / len(lines) if lines else 0
                if pipe_ratio >= 0.15:  # 15% 이상이 테이블 라인
                    return True
        
        # 4. 표 마커와 테이블 구조가 함께 있는 경우
        if re.search(r'\*\*\[표\]?\*\*', content) and len(pipe_lines) >= 2:
            return True
        
        # 5. 테이블 관련 키워드와 구조적 특징 결합 확인
        table_keywords = ['표', '구분', '항목', '분류', '기준', '조건', '업무', '절차', '방법', '내용']
        keyword_found = any(keyword in content for keyword in table_keywords)
        
        if keyword_found and len(pipe_lines) >= 2:
            # 정규 표현식으로 테이블 패턴 감지
            table_pattern = r'\|[^|]*\|[^|]*\|'
            if re.search(table_pattern, content):
                return True
        
        # 6. 구조화된 정보 패턴 감지 (리스트 형태)
        structured_lines = [line for line in lines if re.search(r'[•·▪\-]\s*', line) or ':' in line]
        if len(structured_lines) >= 3 and len(structured_lines) / len(lines) >= 0.25:
            return True
        
        # 7. 페이지 정보와 함께 있는 표 데이터
        if "페이지" in content and len(pipe_lines) >= 2:
            return True
        
        # 8. 표 제목이 있는 경우
        if re.search(r'표\s*\d+', content) and len(pipe_lines) >= 1:
            return True
        
        return False
    
    @staticmethod
    def _extract_text_from_json(json_data: Any, max_depth: int = 10) -> str:
        """JSON 데이터에서 텍스트 추출"""
        if max_depth <= 0:
            return ""
        
        text_parts = []
        
        if isinstance(json_data, dict):
            for key, value in json_data.items():
                text_parts.append(str(key))
                if isinstance(value, (dict, list)):
                    text_parts.append(ElasticsearchIndexer._extract_text_from_json(value, max_depth - 1))
                else:
                    text_parts.append(str(value))
        elif isinstance(json_data, list):
            for item in json_data:
                if isinstance(item, (dict, list)):
                    text_parts.append(ElasticsearchIndexer._extract_text_from_json(item, max_depth - 1))
                else:
                    text_parts.append(str(item))
        else:
            text_parts.append(str(json_data))
        
        return " ".join(filter(None, text_parts))
    
    @staticmethod
    def index_all_files(data_dir: str, embeddings, hybrid_tracker, file_types: List[str] = None, incremental: bool = True) -> Tuple[bool, str]:
        """모든 파일 타입을 한 번에 인덱싱"""
        if file_types is None:
            file_types = ['pdf', 'txt', 'json', 'docx']
        
        hybrid_tracker.track_preprocessing_stage("전체_파일_인덱싱_시작")
        
        # 증분 업데이트 처리
        if not incremental:
            ElasticsearchIndexer._delete_entire_index()
        else:
            categories_to_delete = [ft.upper() for ft in file_types]
            success, message = ElasticsearchIndexer.delete_documents_by_category(categories_to_delete)
            if not success:
                hybrid_tracker.end_preprocessing_stage("전체_파일_인덱싱_시작")
                return False, f"기존 문서 삭제 실패: {message}"
            print(f"📚 {message}")
        
        all_documents = []
        processed_files = 0
        
        # 파일 타입별 처리 매핑
        file_processors = {
            'pdf': (ElasticsearchIndexer.list_pdfs, lambda path: PyPDFLoader(path).load()),
            'txt': (ElasticsearchIndexer.list_txt_files, lambda path: TextLoader(path, encoding='utf-8').load()),
            'json': (ElasticsearchIndexer.list_json_files, ElasticsearchIndexer._process_json_file),
            'docx': (ElasticsearchIndexer.list_docx_files, ElasticsearchIndexer._process_docx_file)
        }
        
        for file_type in file_types:
            if file_type not in file_processors:
                continue
                
            list_func, loader_func = file_processors[file_type]
            type_dir = os.path.join(data_dir, file_type)
            files = list_func(type_dir)
            
            # PDF 파일인 경우 export를 위한 데이터 수집 초기화
            pdf_processing_data = [] if file_type.lower() == 'pdf' else None
            
            for file_path in files:
                stage_name = f"{file_type.upper()}_처리_{os.path.basename(file_path)}"
                hybrid_tracker.track_preprocessing_stage(stage_name)
                
                try:
                    docs = loader_func(file_path)
                    splitter = get_optimized_text_splitter()
                    
                    for doc in docs:
                        # 기본 메타데이터 설정
                        base_metadata = {
                            "source": file_path,
                            "filename": os.path.basename(file_path)
                        }
                        
                        # JSON 파일의 경우 기존 카테고리 정보 보존
                        if file_type.lower() == 'json' and 'category' in doc.metadata:
                            base_metadata["file_type"] = file_type.upper()
                            # 기존 category는 유지
                        else:
                            base_metadata["category"] = file_type.upper()
                        
                        doc.metadata.update(base_metadata)
                        
                        # PDF 파일인 경우 export 데이터 수집
                        if file_type.lower() == 'pdf' and pdf_processing_data is not None:
                            is_table = ElasticsearchIndexer._is_table_content(doc.page_content)
                            pdf_processing_data.append({
                                "content": doc.page_content,
                                "metadata": doc.metadata,
                                "is_table": is_table,
                                "content_length": len(doc.page_content),
                                "filename": os.path.basename(file_path)
                            })
                        
                        if file_type.lower() != 'pdf':
                            # 표가 아닌 경우만 분할
                            if "표" not in doc.page_content:
                                chunks = splitter.split_documents([doc])
                                all_documents.extend(chunks)
                            else:
                                all_documents.append(doc)
                    if file_type.lower() == 'pdf':          ### pdf 경우는 따로 md 테이블 chunker로 chunking
                        chunks = ElasticsearchIndexer.chunk_pdf_with_md_tables(docs,splitter)
                        all_documents.extend(chunks)
                    
                    processed_files += 1
                    
                except Exception as e:
                    print(f"{file_type.upper()} 파일 처리 오류 ({file_path}): {e}")
                finally:
                    hybrid_tracker.end_preprocessing_stage(stage_name)
            
            # PDF 파일 처리 완료 후 export 실행
            if file_type.lower() == 'pdf' and pdf_processing_data:
                try:
                    ElasticsearchIndexer._export_pdf_processing_data(pdf_processing_data, data_dir)
                    print(f"📊 PDF 처리 데이터가 /data 디렉토리에 export되었습니다.")
                except Exception as e:
                    print(f"⚠️ PDF export 중 오류 발생: {e}")
        
        if not all_documents:
            hybrid_tracker.end_preprocessing_stage("전체_파일_인덱싱_시작")
            try:
                es_client, success, message = ElasticsearchManager.get_safe_elasticsearch_client()
                cnt = es_client.count(index=INDEX_NAME).get("count", 0) if success else 0
                return True, f"선택된 카테고리 삭제 완료. 처리된 파일: 0개, 현재 문서 수: {cnt}"
            except:
                return True, "선택된 카테고리 삭제 완료. 처리된 파일: 0개"
        
        # Elasticsearch에 저장
        result, message = ElasticsearchIndexer._save_to_elasticsearch(
            all_documents, embeddings, hybrid_tracker, "전체_파일_인덱싱_시작", "전체"
        )
        
        if result:
            return True, f"전체 파일 인덱싱 완료. 처리된 파일: {processed_files}개, 총 문서 : {message.split('문서 수: ')[1] if '문서 수: ' in message else ''}"
        else:
            return False, message
    
    @staticmethod
    def _process_json_file(path: str) -> List[Document]:
        """JSON 파일 처리 - 검색 최적화를 위한 세분화된 문서 분할"""
        with open(path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        
        documents = []
        
        # converted_index.json 형태의 구조화된 JSON인지 확인
        if isinstance(json_data, list) and len(json_data) > 0:
            for file_entry in json_data:
                if isinstance(file_entry, dict) and 'filename' in file_entry and 'data' in file_entry:
                    # 구조화된 JSON 처리
                    filename = file_entry.get('filename', 'unknown')
                    data_items = file_entry.get('data', [])
                    
                    # 검색 최적화를 위한 다층 구조 처리
                    for item in data_items:
                        if isinstance(item, dict) and item.get('content', '').strip():
                            # 각 항목을 개별 문서로 생성 (최대한 세분화)
                            title = item.get('title', '').strip()
                            heading = item.get('heading', '').strip()
                            section = item.get('section', '').strip()
                            content = item.get('content', '').strip()
                            
                            if not content:
                                continue
                            
                            # 키워드 추출 (메타데이터에서 활용)
                            keywords = ElasticsearchIndexer._extract_keywords(content)
                            
                            # 메타데이터 구성 - 검색 최적화를 위한 풍부한 정보
                            metadata = {
                                'source_filename': filename,
                                'document_type': 'structured_json',
                                'title': title if title else '기타',
                                'heading': heading if heading else '',
                                'section': section if section else '',
                                'category': ElasticsearchIndexer._extract_category(title, heading),
                                'topic': ElasticsearchIndexer._extract_topic(title, heading, section),
                                'content_type': ElasticsearchIndexer._classify_content_type(content),
                                'has_table': item.get('hasTable', False),
                                'keywords': ', '.join(keywords) if keywords else '',  # 검색용 키워드 문자열
                                'keyword_count': len(keywords),  # 키워드 개수
                                'content_length': len(content),  # 내용 길이
                                'hierarchy_depth': len([x for x in [title, heading, section] if x.strip()]),  # 계층 깊이
                                'title_clean': title.strip('[]') if title else '',  # 대괄호 제거한 제목
                                'is_procedure': '절차' in content or '단계' in content,  # 절차 문서 여부
                                'has_numbers': bool(re.search(r'\d+', content)),  # 숫자 포함 여부
                                'urgency_level': 'high' if any(word in content.lower() for word in ['긴급', '즉시', '주의', '경고']) else 'normal'  # 긴급도
                            }
                            
                            # 검색 친화적 문서 내용 구성 - 메타데이터 활용
                            document_parts = []
                            
                            # 제목 계층 구조 추가 (검색 가중치 향상)
                            if title:
                                document_parts.append(f"대분류: {title}")
                                document_parts.append(f"카테고리: {title.strip('[]')}")  # 대괄호 제거한 버전
                            if heading:
                                document_parts.append(f"소분류: {heading}")
                            if section:
                                document_parts.append(f"세부항목: {section}")
                            
                            # 키워드 강화 (검색 정확도 향상)
                            if keywords:
                                document_parts.append(f"주요 키워드: {', '.join(keywords)}")
                                document_parts.append(f"검색어: {' '.join(keywords)}")  # 키워드 반복으로 검색 강화
                            
                            # 문서 특성 표시
                            doc_features = []
                            if item.get('hasTable', False):
                                doc_features.extend(["표 포함", "데이터표", "표형태정보"])
                            if metadata['is_procedure']:
                                doc_features.extend(["절차안내", "단계별가이드"])
                            if metadata['has_numbers']:
                                doc_features.extend(["수치정보", "정량데이터"])
                            if doc_features:
                                document_parts.append(f"문서특성: {', '.join(doc_features)}")
                            
                            # 원본 내용
                            document_parts.append(f"내용: {content}")
                            
                            # 검색용 요약 생성
                            summary = ElasticsearchIndexer._generate_summary(content, title, heading, section)
                            if summary:
                                document_parts.append(f"요약: {summary}")
                            
                            document_content = '\n\n'.join(document_parts)
                            documents.append(Document(page_content=document_content, metadata=metadata))
                    
                    # 대분류별 통합 문서도 생성 (상위 레벨 검색용)
                    title_groups = {}
                    for item in data_items:
                        if isinstance(item, dict):
                            title = item.get('title', '').strip()
                            if title and item.get('content', '').strip():
                                if title not in title_groups:
                                    title_groups[title] = []
                                title_groups[title].append(item)
                    
                    for title, items in title_groups.items():
                        if len(items) > 1:  # 여러 항목이 있는 경우만 통합 문서 생성
                            content_parts = []
                            all_keywords = set()
                            has_any_table = False  # 표 포함 여부 체크
                            
                            for item in items:
                                content = item.get('content', '').strip()
                                if content:
                                    heading = item.get('heading', '').strip()
                                    section = item.get('section', '').strip()
                                    
                                    # 표 포함 여부 확인
                                    if item.get('hasTable', False):
                                        has_any_table = True
                                    
                                    part_header = []
                                    if heading:
                                        part_header.append(heading)
                                    if section:
                                        part_header.append(section)
                                    
                                    if part_header:
                                        content_parts.append(f"[{' - '.join(part_header)}]\n{content}")
                                    else:
                                        content_parts.append(content)
                                    
                                    # 키워드 수집
                                    item_keywords = ElasticsearchIndexer._extract_keywords(content)
                                    all_keywords.update(item_keywords)
                            
                            if content_parts:
                                consolidated_content = '\n\n'.join(content_parts)
                                
                                # 통합 문서의 풍부한 메타데이터
                                metadata = {
                                    'source_filename': filename,
                                    'document_type': 'consolidated_json',
                                    'title': title,
                                    'title_clean': title.strip('[]'),
                                    'category': ElasticsearchIndexer._extract_category(title, ''),
                                    'keywords': ', '.join(sorted(all_keywords)),
                                    'keyword_count': len(all_keywords),
                                    'item_count': len(items),
                                    'has_table': has_any_table,
                                    'content_length': len(consolidated_content),
                                    'is_comprehensive': len(items) > 3,  # 포괄적 문서 여부
                                    'urgency_level': 'high' if any('긴급' in str(item.get('content', '')) for item in items) else 'normal'
                                }
                                
                                # 통합 문서 내용 구성 (검색 최적화)
                                doc_content_parts = [
                                    f"대분류: {title}",
                                    f"카테고리: {title.strip('[]')}",
                                    f"주요 키워드: {', '.join(sorted(all_keywords))}",
                                    f"문서특성: {'표포함 ' if has_any_table else ''}포괄적문서 총{len(items)}개항목",
                                    f"통합 내용:\n{consolidated_content}"
                                ]
                                
                                document_content = '\n\n'.join(doc_content_parts)
                                documents.append(Document(page_content=document_content, metadata=metadata))
                else:
                    # 일반 JSON 처리 (기존 방식)
                    text_content = ElasticsearchIndexer._extract_text_from_json(file_entry)
                    if text_content:
                        documents.append(Document(page_content=text_content, metadata={}))
        else:
            # 일반 JSON 처리 (기존 방식)
            text_content = ElasticsearchIndexer._extract_text_from_json(json_data)
            if text_content:
                documents.append(Document(page_content=text_content, metadata={}))
        
        return documents
    
    @staticmethod
    def _extract_category(title: str, heading: str) -> str:
        """대분류 카테고리 추출 - 개선된 버전"""
        text = f"{title} {heading}".lower()
        
        # 더 정확한 카테고리 분류
        if '기본업무' in text or '처리' in text or '절차' in text:
            return '업무절차'
        elif '소비자' in text or '가이드' in text or '종류' in text:
            return '이용가이드'
        elif '발급' in text or '신청' in text:
            return '카드발급'
        elif '이용한도' in text or '한도' in text:
            return '이용한도'
        elif '연체' in text or '추심' in text or '법적조치' in text:
            return '연체관리'
        elif '부가서비스' in text or '혜택' in text or '포인트' in text:
            return '부가서비스'
        elif '수수료' in text or '연회비' in text or '요금' in text:
            return '수수료정보'
        elif '보안' in text or '분실' in text or '도난' in text:
            return '보안관리'
        elif '해외' in text or '국외' in text:
            return '해외이용'
        elif '현금서비스' in text or '대출' in text or '카드론' in text:
            return '대출서비스'
        else:
            # title에서 카테고리 추출 시도 (대괄호 제거)
            if '[' in title and ']' in title:
                category_match = title.strip('[]').strip()
                if category_match and len(category_match) < 20:
                    return category_match
            return '기타'
    
    @staticmethod
    def _extract_topic(title: str, heading: str, section: str) -> str:
        """세부 주제 추출"""
        parts = [p for p in [title, heading, section] if p.strip()]
        return ' - '.join(parts) if parts else '기타'
    
    @staticmethod
    def _classify_content_type(content: str) -> str:
        """내용 유형 분류"""
        content_lower = content.lower()
        
        if '절차' in content_lower or '단계' in content_lower:
            return '절차안내'
        elif '기준' in content_lower or '조건' in content_lower:
            return '기준정보'
        elif '방법' in content_lower or '활용' in content_lower:
            return '이용방법'
        elif '구분' in content_lower or '종류' in content_lower:
            return '분류정보'
        elif '주의' in content_lower or '금지' in content_lower:
            return '주의사항'
        else:
            return '일반정보'
    
    @staticmethod
    def _extract_keywords(content: str) -> List[str]:
        """내용에서 주요 키워드 추출 - 개선된 버전"""
        import re
        
        # 금융 관련 주요 키워드 패턴 (더 포괄적)
        keyword_patterns = [
            r'신용카드|체크카드|직불카드|선불카드|가족카드|법인카드',
            r'카드발급|발급기준|발급절차|발급조건|발급신청',
            r'이용한도|결제능력|신용등급|신용평점|신용도',
            r'연체|추심|법적조치|연체정보|채권추심',
            r'부가서비스|포인트|할부|적립|혜택|리워드',
            r'연회비|수수료|이자|요금|비용|대금',
            r'VISA|Master|AMEX|JCB|BC카드',
            r'가처분소득|금융거래|신용정보|개인신용평가',
            r'현금서비스|카드론|대출|여신|한도',
            r'분실|도난|재발급|정지|해지|취소',
            r'가맹점|결제|일시불|리볼빙|무이자할부',
            r'해외|국외|국내|온라인|오프라인'
        ]
        
        keywords = set()
        
        # 패턴 매칭으로 키워드 추출
        for pattern in keyword_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            keywords.update(matches)
        
        # 중요한 숫자나 비율 정보
        number_patterns = re.findall(r'\d+(?:\.\d+)?%|\d+만원|\d+개월|\d+일|\d+년', content)
        keywords.update(number_patterns)
        
        # 절차나 단계 관련 키워드
        if '절차' in content or '단계' in content or '방법' in content:
            procedure_words = re.findall(r'[①-⑳]|[가-힣]\)|[1-9]\.|첫째|둘째|셋째', content)
            keywords.update(procedure_words[:5])
        
        # 괄호 안의 중요 정보 추출
        bracket_content = re.findall(r'[「『\[]([^」』\]]+)[」』\]]', content)
        for bc in bracket_content:
            if 2 < len(bc) < 15:  # 적절한 길이의 내용만
                keywords.add(bc)
        
        # 중복 제거 및 정렬, 길이 제한
        result = sorted([k for k in keywords if len(k) > 1 and len(k) < 15])
        return result[:12]  # 최대 12개 키워드
    
    @staticmethod
    def _generate_summary(content: str, title: str, heading: str, section: str) -> str:
        """내용 요약 생성"""
        # 첫 문장이나 핵심 문장 추출
        sentences = content.split('.')
        if sentences:
            first_sentence = sentences[0].strip()
            if len(first_sentence) > 10 and len(first_sentence) < 200:
                return first_sentence
        
        # 제목 정보 기반 요약
        context_parts = [p for p in [title, heading, section] if p.strip()]
        if context_parts:
            return f"{' - '.join(context_parts)}에 대한 정보"
        
        return ""
    
    @staticmethod
    def _process_docx_file(path: str) -> List[Document]:
        """DOCX 파일 처리"""
        try:
            from utils.docparser import DocParser
            parser = DocParser(path)
            text_content = parser.get_text()
            
            docs = [Document(page_content=text_content, metadata={})]
            
            # 표 내용 추가
            table_blocks = parser.get_tables()
            for table_content in table_blocks:
                docs.append(Document(page_content=table_content, metadata={}))
            
            return docs
        except ImportError:
            return ElasticsearchIndexer._docx_basic_loader(path)

